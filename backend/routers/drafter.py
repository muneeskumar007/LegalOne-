"""
routers/drafter.py  — LegalOne Master Drafter Router v4.0
──────────────────────────────────────────────────────────
3-step AI drafting flow with master prompt system:
  POST /api/drafter/step1-extract   → extract facts + detect missing fields
  POST /api/drafter/step2-validate  → validate after user fills gaps
  POST /api/drafter/step3-generate  → generate final court-format petition
  POST /api/drafter/export-docx     → download draft as Word document
  GET  /api/drafter/required-fields/{draft_type}
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
import json
import re
import io
import datetime
import requests

router = APIRouter()

# ─── Ollama config ────────────────────────────────────────────────────────────
OLLAMA_URL   = "http://localhost:11434"
OLLAMA_MODEL = "llama3"

# ─── Field schemas ────────────────────────────────────────────────────────────
REQUIRED_FIELDS: Dict[str, Dict[str, Any]] = {
    "petition": {
        "petitioner_name":       {"label": "Petitioner Name",           "required": True,  "ask": "Full name of the petitioner (wife/husband filing)"},
        "petitioner_age":        {"label": "Petitioner Age",            "required": True,  "ask": "Age of petitioner (in years)"},
        "petitioner_address":    {"label": "Petitioner Address",        "required": True,  "ask": "Current residential address of petitioner"},
        "petitioner_occupation": {"label": "Petitioner Occupation",     "required": False, "ask": "Occupation / profession of petitioner"},
        "respondent_name":       {"label": "Respondent Name",           "required": True,  "ask": "Full name of the respondent (other spouse)"},
        "respondent_age":        {"label": "Respondent Age",            "required": True,  "ask": "Age of respondent (in years)"},
        "respondent_address":    {"label": "Respondent Address",        "required": True,  "ask": "Current residential address of respondent"},
        "marriage_date":         {"label": "Date of Marriage",          "required": True,  "ask": "Date of marriage (DD/MM/YYYY or written out)"},
        "marriage_place":        {"label": "Place of Marriage",         "required": True,  "ask": "City / place where marriage was solemnized"},
        "marriage_type":         {"label": "Marriage Type",             "required": True,  "ask": "Type: Hindu / Muslim / Christian / Special Marriage Act"},
        "marriage_registered":   {"label": "Marriage Registration",     "required": False, "ask": "Was marriage registered? If yes, where?"},
        "cohabitation_address":  {"label": "Last Cohabitation Address", "required": True,  "ask": "Address where parties last lived together as husband & wife"},
        "separation_date":       {"label": "Date of Separation",        "required": True,  "ask": "When did parties separate? (date or approximate year/month)"},
        "children":              {"label": "Children Details",          "required": False, "ask": "Any children? If yes, name, age and sex of each child"},
        "ground":                {"label": "Ground for Divorce",        "required": True,  "ask": "Ground: cruelty / desertion / adultery / mutual consent / others"},
        "cruelty_incidents":     {"label": "Cruelty Incidents",         "required": True,  "ask": "Specific cruelty incidents with dates, nature of acts and effect on petitioner"},
        "jurisdiction_court":    {"label": "Court & City",              "required": True,  "ask": "Which Family Court and city/district to file this petition in"},
        "relief_sought":         {"label": "Relief Sought",             "required": True,  "ask": "Reliefs sought: divorce / + maintenance / + custody / + alimony"},
        "maintenance_amount":    {"label": "Maintenance Amount",        "required": False, "ask": "If maintenance is claimed, how much per month?"},
        "previous_litigation":   {"label": "Previous Court Cases",      "required": False, "ask": "Any previous court cases between the parties? If yes, details"},
        "advocate_name":         {"label": "Advocate Name",             "required": False, "ask": "Name of the advocate filing the petition"},
    },
    "counter": {
        "petitioner_name":  {"label": "Original Petitioner",        "required": True, "ask": "Name of original petitioner / plaintiff"},
        "respondent_name":  {"label": "Respondent / Counter Party", "required": True, "ask": "Name of respondent filing this written statement"},
        "respondent_age":   {"label": "Respondent Age",             "required": True, "ask": "Age of respondent"},
        "respondent_address":{"label": "Respondent Address",        "required": True, "ask": "Current address of respondent"},
        "case_number":      {"label": "Case Number",                "required": True, "ask": "Original case number and court name (e.g. HMA/234/2025, Family Court Delhi)"},
        "reply_points":     {"label": "Reply Points",               "required": True, "ask": "Respondent's version of facts and point-wise reply to each allegation"},
        "relief_sought":    {"label": "Counter Relief Sought",      "required": False, "ask": "Any counter-reliefs: maintenance, custody, return of stridhan etc."},
    },
    "arguments": {
        "petitioner_name": {"label": "Petitioner Name",       "required": True, "ask": "Name of petitioner"},
        "respondent_name": {"label": "Respondent Name",       "required": True, "ask": "Name of respondent"},
        "case_number":     {"label": "Case Number & Court",   "required": True, "ask": "Case number and court (e.g. HMA/234/2025, Family Court Delhi)"},
        "ground":          {"label": "Legal Ground",          "required": True, "ask": "Legal ground / issue (cruelty, desertion, adultery, etc.)"},
        "side":            {"label": "Arguing for",           "required": True, "ask": "Are arguments on behalf of Petitioner or Respondent?"},
        "key_facts":       {"label": "Key Facts for Arguments","required": True, "ask": "Key factual incidents to argue — dates, nature of acts, witnesses"},
        "judgments":       {"label": "Judgments to Cite",     "required": False,"ask": "Any specific judgments to cite? (system will auto-suggest if blank)"},
    },
    # Legacy mapping — keep for backward compatibility
    "divorce_petition": {},
    "legal_notice": {
        "sender_name":      {"label": "Sender Name",     "required": True,  "ask": "Full name of the person sending the notice"},
        "receiver_name":    {"label": "Receiver Name",   "required": True,  "ask": "Full name of the person receiving the notice"},
        "receiver_address": {"label": "Receiver Address","required": True,  "ask": "Full postal address of the receiver"},
        "subject":          {"label": "Subject / Demand","required": True,  "ask": "What is the notice about? What is being demanded?"},
        "amount":           {"label": "Amount (if any)", "required": False, "ask": "Amount of money involved, if any (with currency)"},
        "deadline_days":    {"label": "Reply Deadline",  "required": True,  "ask": "Days given to respond (usually 15 or 30 days)"},
        "facts":            {"label": "Facts",           "required": True,  "ask": "Detailed facts — what happened, when, and why this notice is sent"},
    },
    "counter_reply": {
        "petitioner_name": {"label": "Original Petitioner",        "required": True, "ask": "Name of original petitioner"},
        "respondent_name": {"label": "Respondent / Counter Party",  "required": True, "ask": "Name of respondent filing this counter reply"},
        "case_number":     {"label": "Case Number",                 "required": True, "ask": "Original case number and court name"},
        "reply_points":    {"label": "Reply Points",                "required": True, "ask": "Point-wise reply to each allegation in the petition"},
        "defence":         {"label": "Affirmative Defence",         "required": True, "ask": "Respondent's defence and any counter-claims"},
    },
}


# ─── Pydantic models ──────────────────────────────────────────────────────────

class Step1Request(BaseModel):
    draft_type:  str
    user_input:  str
    rag_context: Optional[str] = ""

class Step2Request(BaseModel):
    draft_type:       str
    original_input:   str
    user_responses:   Dict[str, Any]
    extracted_so_far: Dict[str, Any]

class Step3Request(BaseModel):
    draft_type:   str
    final_data:   Dict[str, Any]
    rag_context:  Optional[str] = ""
    side:         Optional[str] = "petitioner"   # for arguments
    petition_text:Optional[str] = ""             # for counter

class DocxExportRequest(BaseModel):
    draft_text: str
    filename:   Optional[str] = "LegalOne_Draft"


# ─── Ollama caller ────────────────────────────────────────────────────────────

def _call_ollama(system: str, user: str, timeout: int = 180, max_tokens: int = 4000) -> Optional[str]:
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/chat",
            json={
                "model": OLLAMA_MODEL,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user},
                ],
                "stream": False,
                "options": {"temperature": 0.05, "num_predict": max_tokens},
            },
            timeout=timeout,
        )
        if resp.status_code == 200:
            return resp.json().get("message", {}).get("content", "")
    except Exception as e:
        print(f"[Drafter] Ollama error: {e}")
    return None


def _parse_json(raw: str) -> dict:
    """Strip markdown fences and parse first JSON object found."""
    raw = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`").strip()
    start = raw.find("{")
    end   = raw.rfind("}")
    if start != -1 and end != -1:
        try:
            return json.loads(raw[start : end + 1])
        except Exception:
            pass
    return {}


def _get_rag_context(query: str, top_k: int = 4) -> str:
    try:
        from core.rag_pipeline import get_context_for_prompt
        return get_context_for_prompt(query, top_k=top_k)
    except Exception:
        return ""


def _normalize_type(draft_type: str) -> str:
    """Map legacy API types to new canonical types."""
    mapping = {
        "divorce_petition": "petition",
        "counter_reply":    "counter",
        "counter-reply":    "counter",
        "legal_notice":     "legal_notice",
    }
    return mapping.get(draft_type, draft_type)


# ═══════════════════════════════════════════════════════════════════════════════
# MASTER PROMPT SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

BASE_SYSTEM = """You are a Senior Indian Advocate with 25+ years of experience drafting
matrimonial, civil, and criminal pleadings in Indian courts.

ABSOLUTE RULES — never break these:
1. Use ONLY the facts given. Never invent dates, names, addresses, or incidents.
2. Where data is missing, write [TO BE FILLED] — never guess.
3. Follow the EXACT formatting rules specified — bold markers, spacing, numbering.
4. Output ONLY the draft text. No explanation, no preamble, no markdown fences.
5. Use formal Indian legal English throughout.
6. Every paragraph must be numbered sequentially.
7. Prayer clause items must be lettered (a), (b), (c)...
8. Verification paragraph must be present and correct.
9. Never add facts that were not provided."""


PETITION_SYSTEM = BASE_SYSTEM + """

EXACT FORMAT RULES FOR DIVORCE PETITION:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
LINE 1  (BOLD CAPS): IN THE COURT OF THE [COURT NAME] AT [CITY]
BLANK LINE
LINE 2  (BOLD CAPS): MATRIMONIAL CASE NO. __________ OF 20__
BLANK LINE
IN THE MATTER OF:
BLANK LINE
[Petitioner Name], [s/o or d/o or w/o] [Father/Husband name],
Aged about [age] years, [Occupation],
Residing at [Full Address].
                                                           ...PETITIONER
BLANK LINE
                        VERSUS
BLANK LINE
[Respondent Name], [s/o or d/o or w/o] [Father/Husband name],
Aged about [age] years, [Occupation],
Residing at [Full Address].
                                                          ...RESPONDENT
BLANK LINE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): PETITION FOR DECREE OF DIVORCE UNDER SECTION 13(1)(ia) OF THE HINDU MARRIAGE ACT, 1955
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BLANK LINE
MOST RESPECTFULLY SHOWETH:
BLANK LINE
The Petitioner above named respectfully states as follows:
BLANK LINE
1.  [Marriage paragraph — date, place, rites, registration]
BLANK LINE
2.  [Status paragraph — addresses before marriage and now]
BLANK LINE
3.  [Children paragraph — or "No child has been born out of the wedlock"]
BLANK LINE
4.  [Cohabitation and start of problems paragraph]
BLANK LINE
5 onwards — [ONE paragraph per cruelty incident: date → nature → effect]
BLANK LINE
[N].   [Condonation paragraph]
[N+1]. [Collusion paragraph]
[N+2]. [Limitation paragraph]
[N+3]. [Previous litigation paragraph]
[N+4]. [Jurisdiction paragraph]
[N+5]. [Court fees paragraph]
BLANK LINE
(BOLD CAPS): PRAYER
BLANK LINE
That this Hon'ble Court be pleased to:
(a)  Pass a decree of divorce...
(b)  [additional reliefs]
(c)  Pass such other and further orders...
BLANK LINE
                                          Respectfully submitted,
BLANK LINE
Place: [City]                             [Petitioner Name]
Date: [Date]                              PETITIONER
BLANK LINE
Through:
[Advocate Name]
Advocate
BLANK LINE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): VERIFICATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BLANK LINE
I, [Petitioner Name], the Petitioner above named, do hereby verify that the
contents of paragraphs 1 to [N] are true to my personal knowledge and belief,
and the contents of paragraphs [N+1] to [total] are true to the best of my
information and belief.
BLANK LINE
Verified at [City] on this _____ day of _________ 20__.
BLANK LINE
                                          [Petitioner Name]
                                          PETITIONER"""


COUNTER_SYSTEM = BASE_SYSTEM + """

EXACT FORMAT RULES FOR WRITTEN STATEMENT / COUNTER:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): IN THE COURT OF THE [COURT NAME] AT [CITY]
(BOLD CAPS): MATRIMONIAL CASE NO. [NUMBER] OF 20__
BLANK LINE
IN THE MATTER OF:
[Petitioner] ...PETITIONER
                    VERSUS
[Respondent] ...RESPONDENT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): WRITTEN STATEMENT / COUNTER FILED BY THE RESPONDENT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BLANK LINE
(BOLD CAPS): PRELIMINARY OBJECTIONS:
BLANK LINE
I.   [First preliminary objection — maintainability]
II.  [Second preliminary objection — jurisdiction]
III. [Third preliminary objection — limitation or suppression of facts]
BLANK LINE
(BOLD CAPS): WITHOUT PREJUDICE TO THE ABOVE OBJECTIONS, THE RESPONDENT STATES ON MERITS:
BLANK LINE
(BOLD CAPS): CASE FACTS:
BLANK LINE
1.  [Respondent's own version of the facts]
2+. [Continue facts paragraphs]
BLANK LINE
(BOLD CAPS): PARAWISE REPLY:
BLANK LINE
Para 1 of the Petition is [admitted/denied]. [Brief reply]
Para 2 of the Petition is [admitted/denied]. [Brief reply]
[Continue for each para]
BLANK LINE
(BOLD CAPS): REPLY TO PRAYER:
BLANK LINE
The Respondent denies that the Petitioner is entitled to any of the reliefs...
BLANK LINE
(BOLD CAPS): PRAYER:
BLANK LINE
(a) Dismiss the petition with costs;
(b) [Any counter-relief]
(c) Pass any other order as deemed fit.
BLANK LINE
                                         Respectfully submitted,
Place: [City]                            [Respondent Name]
Date:  [Date]                            RESPONDENT
BLANK LINE
Through:
[Advocate Name]
Advocate for Respondent
BLANK LINE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): VERIFICATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
I, [Respondent Name], do hereby verify that the contents of this Written Statement
are true to my knowledge and belief...
Verified at [City] on this _____ day of _________ 20__."""


ARGUMENTS_SYSTEM = BASE_SYSTEM + """

EXACT FORMAT FOR WRITTEN ARGUMENTS:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): IN THE COURT OF THE [COURT] AT [CITY]
(BOLD CAPS): MATRIMONIAL CASE NO. [NUMBER] OF 20__
[PETITIONER] vs [RESPONDENT]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
(BOLD CAPS): WRITTEN ARGUMENTS ON BEHALF OF [PETITIONER/RESPONDENT]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BLANK LINE
(BOLD): A. BRIEF FACTS
BLANK LINE
[2-3 paragraph factual summary from the arguing party's perspective]
BLANK LINE
(BOLD): B. ISSUES INVOLVED
BLANK LINE
Issue No. 1: [Issue heading]
Issue No. 2: [Issue heading]
BLANK LINE
(BOLD): C. ARGUMENTS
BLANK LINE
(BOLD): Issue No.1: [Issue heading]
BLANK LINE
1.  [Legal argument with citation]
2.  [Next argument — cite sections + judgments]
BLANK LINE
(BOLD): Issue No.2: [Issue heading]
BLANK LINE
[Continue per issue]
BLANK LINE
(BOLD): D. CONCLUSION
BLANK LINE
In light of the above submissions and judgments, it is most humbly submitted
that this Hon'ble Court be pleased to...
BLANK LINE
                                         Respectfully submitted,
Place: [City]
Date:  [Date]                            [Party Name]
                                         Through [Advocate]"""


def _petition_prompt(facts: dict, rag_context: str = "") -> dict:
    data = "\n".join(
        f"  {k}: {v}" for k, v in facts.items()
        if v and v not in ([], None, "", "null")
    )
    user = f"""GENERATE DIVORCE PETITION — SECTION 13(1)(ia) HMA

COMPLETE CASE DATA:
{data}

{f"RAG LEGAL CONTEXT (from real judgments):{chr(10)}{rag_context[:1200]}" if rag_context else ""}

DRAFTING INSTRUCTIONS:
1. Para 1 → Marriage details (date, place, rites; mention Annexure P-1 for certificate)
2. Para 2 → Status of parties (addresses at time of marriage + current)
3. Para 3 → Children (use children data; or state "No child has been born")
4. Para 4 → Cohabitation and when problems started
5. Paras 5+ → ONE paragraph per cruelty incident:
   - State approximate date
   - State exact nature of act
   - State effect on petitioner's physical/mental health
   - Use: "treated the Petitioner with cruelty"
6. After cruelty: Condonation → Collusion → Limitation → Previous Litigation → Jurisdiction → Court Fees
7. PRAYER — (a), (b), (c)... list each relief
8. VERIFICATION — state which paras are own knowledge vs information

Start directly with "IN THE COURT OF..." """

    return {"system": PETITION_SYSTEM, "user": user}


def _counter_prompt(facts: dict, petition_text: str = "", rag_context: str = "") -> dict:
    data = "\n".join(f"  {k}: {v}" for k, v in facts.items() if v and v not in ([], None))
    user = f"""GENERATE WRITTEN STATEMENT / COUNTER PETITION

RESPONDENT'S CASE DATA:
{data}

{f"ORIGINAL PETITION TEXT (to reply para-wise):{chr(10)}{petition_text[:2000]}" if petition_text else ""}

{f"RAG LEGAL CONTEXT:{chr(10)}{rag_context[:800]}" if rag_context else ""}

DRAFTING INSTRUCTIONS:
1. PRELIMINARY OBJECTIONS — 3-4 (maintainability, limitation, jurisdiction, suppression)
2. CASE FACTS — respondent's version (deny cruelty; state respondent's conduct was justified)
3. PARAWISE REPLY — reply to each numbered petition para:
   - Para 1: Usually admitted (marriage facts)
   - Para 2: Admitted in part (addresses)
   - Para 3: Admitted or denied (children)
   - Para 4+: Denied — specific denial + respondent's counter version
4. REPLY TO PRAYER — deny all reliefs sought
5. PRAYER — dismiss with costs + any counter-reliefs (maintenance, custody, stridhan)
6. VERIFICATION

Start directly with "IN THE COURT OF..." """

    return {"system": COUNTER_SYSTEM, "user": user}


def _arguments_prompt(facts: dict, side: str, rag_context: str = "") -> dict:
    data = "\n".join(f"  {k}: {v}" for k, v in facts.items() if v and v not in ([], None))
    user = f"""GENERATE WRITTEN ARGUMENTS FOR {side.upper()}

CASE DATA:
{data}

SIDE: {side} (Petitioner = arguing FOR divorce; Respondent = arguing AGAINST)

LEGAL CONTEXT FROM REAL JUDGMENT DATASET:
{rag_context or "[No RAG context — use standard matrimonial law principles and Supreme Court judgments]"}

DRAFTING INSTRUCTIONS:
1. BRIEF FACTS — 2-3 para summary from {side}'s perspective
2. ISSUES — frame 2-4 issues (cruelty proved/not, desertion, jurisdiction etc.)
3. ARGUMENTS — per issue:
   - State legal position
   - Cite relevant sections (HMA, CrPC, Evidence Act)
   - Cite 2-3 real judgments from RAG context if available
   - Apply law to facts of THIS case
   - For cruelty: cite Samar Ghosh v. Jaya Ghosh (2007) 4 SCC 511 — 31 guiding principles
   - For desertion: cite Bipinchandra Jaisinghbhai Shah v. Prabhavati (1957 AIR SC)
   - For false complaint: cite V. Bhagat v. D. Bhagat (1994) 1 SCC 337
4. CONCLUSION — prayer based on arguments
5. Use strong, persuasive language appropriate for court submissions

Start directly with "IN THE COURT OF..." """

    return {"system": ARGUMENTS_SYSTEM, "user": user}


# ─── FACT SATURATION SYSTEM ───────────────────────────────────────────────────

FACT_SATURATION_SYSTEM = """You are a senior Indian legal assistant checking if case details
are complete enough to draft a court petition.

You ONLY return a JSON object. No explanation. No markdown. Pure JSON.

Your job:
1. Extract everything already provided from the user's text
2. List what is MISSING and REQUIRED to draft this specific petition type
3. Be specific about what's missing — "cruelty incidents" means dates, nature, effect
4. Return completeness_score 0-100
5. If score >= 80 and all required fields present, set ready_to_draft: true"""


# ─── STEP 1 — Extraction + Missing Field Detection ───────────────────────────

def _ollama_extract(draft_type: str, user_input: str, rag_context: str, extracted_so_far: dict) -> dict:
    """Use LLM to detect missing fields. Start with regex-extracted data."""
    schema = REQUIRED_FIELDS.get(draft_type, REQUIRED_FIELDS["petition"])
    required_list = "\n".join(
        f"  - {k}: {v['ask']} [{'REQUIRED' if v['required'] else 'optional'}]"
        for k, v in schema.items()
    )

    # Build a string showing what regex already extracted
    already = "\n".join(
        f"  {k}: {v}" for k, v in extracted_so_far.items()
        if v and v not in ([], None, "", "null")
    ) or "  (nothing extracted yet)"

    user = f"""DRAFT TYPE: {draft_type.upper()}

USER PROVIDED TEXT:
{user_input}

ALREADY EXTRACTED BY REGEX:
{already}

{f"RELEVANT LEGAL CONTEXT (RAG):{chr(10)}{rag_context}" if rag_context else ""}

REQUIRED FIELDS FOR {draft_type.upper()}:
{required_list}

Check what is provided vs what is required. Merge regex extraction with your extraction.
Return ONLY this JSON (no explanation):
{{
  "extracted": {{
    "petitioner_name": "value or null",
    "petitioner_age": "value or null",
    "petitioner_address": "value or null",
    "petitioner_occupation": "value or null",
    "respondent_name": "value or null",
    "respondent_age": "value or null",
    "respondent_address": "value or null",
    "respondent_occupation": "value or null",
    "marriage_date": "value or null",
    "marriage_place": "value or null",
    "marriage_type": "value or null",
    "separation_date": "value or null",
    "cohabitation_address": "value or null",
    "children": "value or null",
    "ground": "value or null",
    "cruelty_incidents": "value or null",
    "jurisdiction_court": "value or null",
    "jurisdiction_city": "value or null",
    "relief_sought": "value or null",
    "maintenance_amount": "value or null",
    "previous_litigation": "value or null",
    "advocate_name": "value or null"
  }},
  "missing_required": [
    {{"field": "field_key", "label": "Human readable label", "question": "Exact question to show advocate", "why_needed": "One line legal reason"}}
  ],
  "missing_optional": [
    {{"field": "field_key", "label": "Label", "question": "Question text"}}
  ],
  "completeness_score": 0,
  "ready_to_draft": false,
  "issues": []
}}"""

    raw    = _call_ollama(FACT_SATURATION_SYSTEM, user, timeout=120, max_tokens=2000)
    result = _parse_json(raw) if raw else {}
    return result


def _fallback_extract(draft_type: str, user_input: str, regex_extracted: dict) -> dict:
    """Rule-based extraction when Ollama is unavailable — augments regex extraction."""
    schema     = REQUIRED_FIELDS.get(draft_type, REQUIRED_FIELDS["petition"])
    extracted  = dict(regex_extracted)  # start from regex results

    text     = user_input
    text_low = text.lower()

    # Fill in fields not captured by regex
    if not extracted.get("ground"):
        if "mutual consent" in text_low:
            extracted["ground"] = "mutual consent"
        elif "cruelty" in text_low:
            extracted["ground"] = "cruelty"
        elif "desertion" in text_low:
            extracted["ground"] = "desertion"
        elif "adultery" in text_low:
            extracted["ground"] = "adultery"

    if not extracted.get("relief_sought"):
        if "maintenance" in text_low and "divorce" in text_low:
            extracted["relief_sought"] = "Divorce decree and monthly maintenance"
        elif "custody" in text_low and "divorce" in text_low:
            extracted["relief_sought"] = "Divorce decree and custody of children"
        elif "divorce" in text_low:
            extracted["relief_sought"] = "Divorce decree"

    if not extracted.get("cruelty_incidents"):
        # Look for numbered incidents
        bullets = re.findall(r"\d+\.\s+(?:On\s+\d|That\s+the\s+respondent)[^\n]{20,}", text)
        if bullets:
            extracted["cruelty_incidents"] = "\n".join(bullets[:8])

    if not extracted.get("jurisdiction_court"):
        m = re.search(r"(?:at|before|in)\s+the\s+(Family\s+Court[^,.\n]+|District\s+Court[^,.\n]+)", text, re.I)
        if m:
            extracted["jurisdiction_court"] = m.group(1).strip()

    # Build missing lists
    missing_required: List[dict] = []
    missing_optional: List[dict] = []

    for k, v in schema.items():
        val = extracted.get(k)
        if val and str(val).strip() and str(val).strip() not in ("null", "none", "[]"):
            continue
        entry = {"field": k, "label": v["label"], "question": v["ask"]}
        if v["required"]:
            entry["why_needed"] = "Required to draft the petition correctly"
            missing_required.append(entry)
        else:
            missing_optional.append(entry)

    filled   = sum(1 for k in schema if extracted.get(k) and str(extracted.get(k, "")).strip())
    total    = len(schema) or 1
    score    = min(100, int(filled / total * 100))
    complete = len(missing_required) == 0

    return {
        "extracted":          extracted,
        "missing_required":   missing_required,
        "missing_optional":   missing_optional,
        "is_complete":        complete,
        "completeness_score": score,
        "ready_to_draft":     complete,
    }


# ─── STEP 2 — Validator ───────────────────────────────────────────────────────

def _ollama_validate(draft_type: str, original_input: str,
                     user_responses: dict, extracted: dict) -> dict:
    merged      = {**extracted, **{k: v for k, v in user_responses.items() if v}}
    merged_text = "\n".join(f"  {k}: {v}" for k, v in merged.items())

    system = """You are a senior Indian legal assistant validating case information before drafting.
Check all provided information and decide:
1. Is every REQUIRED field now filled with SPECIFIC information?
2. Are cruelty incidents described with enough detail for court?
3. Is jurisdiction established? Are parties fully identified?

Respond ONLY with valid JSON:
{
  "still_missing": [
    { "field": "field_name", "label": "Label", "question": "What is still needed?", "severity": "critical" }
  ],
  "all_complete": true,
  "completeness_score": 95,
  "issues_found": [],
  "ready_to_draft": true,
  "final_data": { "petitioner_name": "...", ... all merged fields ... }
}"""

    user_msg = f"""DRAFT TYPE: {draft_type}

ORIGINAL INPUT:
{original_input}

ALL DATA MERGED:
{merged_text}

Validate completeness and return JSON with final_data containing ALL fields merged."""

    raw    = _call_ollama(system, user_msg, timeout=120, max_tokens=2000)
    result = _parse_json(raw) if raw else {}
    return result


def _fallback_validate(draft_type: str, user_responses: dict, extracted: dict) -> dict:
    schema = REQUIRED_FIELDS.get(draft_type, REQUIRED_FIELDS["petition"])
    merged = {**extracted, **{k: v for k, v in user_responses.items() if v and str(v).strip()}}

    still_missing = []
    for k, v in schema.items():
        if v["required"] and not str(merged.get(k, "")).strip():
            still_missing.append({
                "field": k, "label": v["label"],
                "question": v["ask"], "severity": "critical",
            })

    filled   = sum(1 for k in schema if str(merged.get(k, "")).strip())
    total    = len(schema) or 1
    score    = int(filled / total * 100)
    complete = len(still_missing) == 0

    return {
        "still_missing":      still_missing,
        "all_complete":       complete,
        "completeness_score": score,
        "issues_found":       [],
        "ready_to_draft":     complete,
        "final_data":         merged,
    }


# ─── STEP 3 — Draft Generator ─────────────────────────────────────────────────

def _ollama_generate(draft_type: str, final_data: dict, rag_context: str,
                     side: str = "petitioner", petition_text: str = "") -> str:
    """Route to correct prompt based on draft type."""
    if draft_type in ("counter", "counter_reply", "counter-reply"):
        prompts = _counter_prompt(final_data, petition_text, rag_context)
        tokens  = 4000
    elif draft_type == "arguments":
        prompts = _arguments_prompt(final_data, side, rag_context)
        tokens  = 3500
    else:
        # petition / divorce_petition / legal_notice / proof-affidavit / other
        prompts = _petition_prompt(final_data, rag_context)
        tokens  = 4000

    return _call_ollama(prompts["system"], prompts["user"], timeout=200, max_tokens=tokens) or ""


def _template_generate(draft_type: str, final_data: dict) -> str:
    """
    Template-based fallback draft when Ollama is offline.
    Produces a complete, properly structured petition.
    """
    year = datetime.datetime.now().year
    d = final_data

    pname  = d.get("petitioner_name",    "The Petitioner")
    page   = d.get("petitioner_age",     "___")
    paddr  = d.get("petitioner_address", "______________________________________")
    pocc   = d.get("petitioner_occupation", "")
    rname  = d.get("respondent_name",    "The Respondent")
    rage   = d.get("respondent_age",     "___")
    raddr  = d.get("respondent_address", "______________________________________")
    mdate  = d.get("marriage_date",      "__________")
    mplace = d.get("marriage_place",     "__________")
    mtype  = d.get("marriage_type",      "Hindu")
    cohab  = d.get("cohabitation_address", raddr)
    sep    = d.get("separation_date",    "__________")
    ground = d.get("ground",             "cruelty").lower()
    court  = d.get("jurisdiction_court", d.get("court", "____________"))
    city   = d.get("jurisdiction_city",  d.get("location", mplace))
    relief = d.get("relief_sought",      "Divorce decree")
    cruelty_raw = d.get("cruelty_incidents", "")
    if isinstance(cruelty_raw, list):
        cruelty_raw = "\n".join(cruelty_raw)
    children = d.get("children", "")
    if isinstance(children, list):
        children = ", ".join(
            c.get("name", "") + (f" (aged {c['age']} yrs)" if c.get("age") else "")
            for c in children if c
        )
    advocate     = d.get("advocate_name", "Advocate")
    prev_lit     = d.get("previous_litigation", "")
    maintenance  = d.get("maintenance_amount", "")
    pocc_str     = f", {pocc}" if pocc else ""

    # Children paragraph
    if children and children.strip().lower() not in ("none", "no", "nil", ""):
        children_para = f"That out of the wedlock, the following child/children were born: {children}."
    else:
        children_para = "That no child was born out of the said wedlock."

    # Cruelty paragraphs
    cruelty_paras = ""
    if cruelty_raw:
        incidents = [x.strip() for x in re.split(r"[.\n]", cruelty_raw) if x.strip() and len(x.strip()) > 10]
        for i, inc in enumerate(incidents[:6], 5):
            cruelty_paras += f"\n{i}.  That {inc}.\n"
    else:
        cruelty_paras = f"\n5.  That the Respondent treated the Petitioner with cruelty during the subsistence of the marriage. The Petitioner was subjected to mental and physical cruelty on numerous occasions.\n"

    next_para = 5 + max(1, len([x for x in re.split(r"[.\n]", cruelty_raw) if x.strip()]) if cruelty_raw else 1)

    # Prayer
    prayer_lines = [f"(a) Pass a decree of divorce dissolving the marriage between {pname} and {rname};"]
    if maintenance:
        prayer_lines.append(f"(b) Direct the Respondent to pay maintenance of Rs. {maintenance}/- per month to the Petitioner;")
    if "custody" in (relief or "").lower():
        prayer_lines.append(f"({'bc'[len(prayer_lines)-1]}) Grant custody of the minor child/children to the Petitioner;")
    prayer_lines.append(f"({'abcde'[len(prayer_lines)]}) Award costs of this petition to the Petitioner;")
    prayer_lines.append(f"({'abcde'[len(prayer_lines)]}) Grant any other relief as this Hon'ble Court may deem fit and proper.")
    prayer_text = "\n".join(prayer_lines)

    prev_para = ""
    if prev_lit and prev_lit.strip().lower() not in ("none", "no", "nil", ""):
        prev_para = f"\n{next_para}.  That the following earlier proceedings have been initiated between the parties: {prev_lit}.\n"
        next_para += 1

    jx_ground = (f"the marriage was solemnized at {mplace}" if mplace != "__________"
                 else "the parties last resided together within the jurisdiction of this Court")

    draft = f"""IN THE COURT OF THE FAMILY JUDGE AT {court.upper() if court else "[CITY]"}

MATRIMONIAL CASE NO. __________ OF {year}


IN THE MATTER OF:

{pname},
Aged about {page} years{pocc_str},
Residing at {paddr}.
                                                      ...PETITIONER

                         VERSUS

{rname},
Aged about {rage} years,
Residing at {raddr}.
                                                      ...RESPONDENT


PETITION FOR DECREE OF DIVORCE UNDER SECTION 13(1)(ia)
OF THE HINDU MARRIAGE ACT, 1955


MOST RESPECTFULLY SHOWETH:

The Petitioner above named states as under:

1.  That the marriage of the Petitioner and the Respondent was
    solemnized on {mdate} at {mplace} according to {mtype} rites and
    ceremonies. {f"The marriage was duly registered." if d.get("marriage_registered") else "A copy of the marriage certificate is annexed hereto as Annexure P-1."}

2.  That the status and place of residence of the parties at the
    time of filing this petition is as under:

    (i)  Petitioner: {paddr}
    (ii) Respondent: {raddr}

3.  {children_para}

4.  That after the solemnization of the marriage, the parties
    resided together at {cohab}. The Respondent, however, started
    treating the Petitioner with cruelty. The details of such
    cruelty are set out hereinbelow:
{cruelty_paras}
{next_para}.  That the Petitioner has not condoned any of the acts of
    cruelty by the Respondent. It has become impossible for the
    Petitioner to live with the Respondent. The parties have been
    living separately since {sep}.

{next_para + 1}.  That there is no collusion between the Petitioner and the
    Respondent in filing this petition. The petition is not presented
    in collusion with the Respondent.

{next_para + 2}.  That the petition is filed within the period of limitation
    as prescribed under the Hindu Marriage Act, 1955. The cause of
    action is continuing.
{prev_para}
{next_para + 3}.  That this Hon'ble Court has jurisdiction to entertain this
    petition as {jx_ground} and falls within the territorial
    jurisdiction of this Court.

{next_para + 4}.  That the requisite court fee has been paid / will be paid as
    per the applicable schedule.


                                PRAYER

It is therefore most humbly prayed that this Hon'ble Court may be
pleased to:

{prayer_text}


                                         Respectfully submitted,

                                         {pname}
                                         PETITIONER

Through:

{advocate}
Advocate

Place: {city}
Date : __________


━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                           VERIFICATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

I, {pname}, aged {page} years, residing at {paddr}, the Petitioner
above named, do hereby verify and declare that the contents of
paragraphs 1 to {next_para + 2} of this petition are true and correct
to the best of my knowledge and belief, and that paragraphs
{next_para + 3} to {next_para + 4} are based on legal advice received
and are believed to be true.

Verified at {city} on this _______ day of _______ {year}.


                                         {pname}
                                         PETITIONER

[NOTE: This draft was generated using the rule-based template engine
as the Ollama AI model is currently offline. For a more precise
AI-generated draft, please ensure Ollama is running with: ollama run llama3]
"""
    return draft


# ─── DOCX EXPORT ─────────────────────────────────────────────────────────────

def _build_docx(draft_text: str) -> bytes:
    """
    Convert plain-text draft to a properly formatted Word document.
    - ALL-CAPS lines → bold
    - Lines starting with ━ → horizontal rule style
    - Normal lines → normal style
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    lines = draft_text.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip separator lines
        if set(stripped) <= {"━", "─", "-", "=", " "} and len(stripped) > 3:
            p = doc.add_paragraph()
            p.add_run("─" * 70)
            continue

        # Empty line → blank paragraph
        if not stripped:
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)

        # Detect formatting cues
        is_bold = False
        align   = WD_ALIGN_PARAGRAPH.LEFT

        # ALL-CAPS lines (court header, section headers, prayer, verification)
        if (stripped == stripped.upper() and len(stripped) > 5
                and any(c.isalpha() for c in stripped)):
            is_bold = True
            align   = WD_ALIGN_PARAGRAPH.CENTER if len(stripped) < 80 else WD_ALIGN_PARAGRAPH.LEFT

        # Lines starting with ...PETITIONER / ...RESPONDENT
        if stripped.startswith("..."):
            align   = WD_ALIGN_PARAGRAPH.RIGHT
            is_bold = True

        # VERSUS
        if stripped == "VERSUS":
            align   = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True

        # PRAYER / VERIFICATION labels
        if stripped in ("PRAYER", "VERIFICATION", "MOST RESPECTFULLY SHOWETH:"):
            is_bold = True
            align   = WD_ALIGN_PARAGRAPH.CENTER

        # Preserve indentation for numbered paragraphs
        leading_spaces = len(line) - len(line.lstrip())
        if leading_spaces >= 4:
            p.paragraph_format.left_indent = Cm(1.0)

        p.alignment = align
        run = p.add_run(stripped)
        run.bold = is_bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── API Routes ───────────────────────────────────────────────────────────────

@router.post("/drafter/step1-extract")
def step1_extract(req: Step1Request):
    """
    Step 1: Extract facts from user input and detect missing fields.
    Called when user clicks 'Generate Draft'.
    """
    if not req.user_input.strip():
        raise HTTPException(400, "Please describe your case first")

    draft_type = _normalize_type(req.draft_type)

    # Step 1a: Regex extraction (fast, reliable for named fields)
    from core.fact_extractor import extract_facts
    regex_result = extract_facts(req.user_input)

    # Map regex result keys to schema keys
    regex_extracted = {
        "petitioner_name":       regex_result.get("petitioner_name"),
        "petitioner_age":        regex_result.get("petitioner_age"),
        "petitioner_address":    regex_result.get("petitioner_address"),
        "petitioner_occupation": regex_result.get("petitioner_occupation"),
        "respondent_name":       regex_result.get("respondent_name"),
        "respondent_age":        regex_result.get("respondent_age"),
        "respondent_address":    regex_result.get("respondent_address"),
        "respondent_occupation": regex_result.get("respondent_occupation"),
        "marriage_date":         regex_result.get("marriage_date"),
        "marriage_place":        regex_result.get("marriage_place"),
        "marriage_type":         regex_result.get("marriage_type"),
        "separation_date":       regex_result.get("separation_date"),
        "cohabitation_address":  regex_result.get("cohabitation_address"),
        "children":              str(regex_result.get("children", "") or ""),
        "ground":                regex_result.get("ground"),
        "cruelty_incidents":     str(regex_result.get("cruelty_incidents", "") or ""),
        "jurisdiction_court":    regex_result.get("jurisdiction_court") or regex_result.get("court"),
        "jurisdiction_city":     regex_result.get("jurisdiction_city") or regex_result.get("location"),
        "relief_sought":         regex_result.get("relief_sought") or regex_result.get("relief"),
        "maintenance_amount":    regex_result.get("maintenance_amount") or regex_result.get("amount"),
        "previous_litigation":   regex_result.get("previous_litigation"),
        "advocate_name":         regex_result.get("advocate_name"),
    }
    # Remove None/empty
    regex_extracted = {k: v for k, v in regex_extracted.items() if v and str(v).strip() and str(v).strip() not in ("null", "None", "[]")}

    # Step 1b: LLM extraction to find missing fields
    rag_ctx = req.rag_context or _get_rag_context(req.user_input, top_k=4)
    result  = _ollama_extract(draft_type, req.user_input, rag_ctx, regex_extracted)

    if not result or "extracted" not in result:
        # Fallback: rule-based
        result = _fallback_extract(draft_type, req.user_input, regex_extracted)
        result["source"] = "rule_based"
    else:
        # Merge: regex takes priority for fields it found, LLM fills the rest
        llm_extracted = result.get("extracted", {})
        for k, v in llm_extracted.items():
            if v and v not in ("null", "none", "", "None") and k not in regex_extracted:
                regex_extracted[k] = v
        result["extracted"] = regex_extracted
        result["source"]    = "ollama"

    return {
        "success":            True,
        "extracted":          result.get("extracted", {}),
        "missing_required":   result.get("missing_required", []),
        "missing_optional":   result.get("missing_optional", []),
        "completeness_score": result.get("completeness_score", 0),
        "ready_to_draft":     result.get("ready_to_draft", False),
        "show_popup":         len(result.get("missing_required", [])) > 0,
        "source":             result.get("source", "unknown"),
    }


@router.post("/drafter/step2-validate")
def step2_validate(req: Step2Request):
    """
    Step 2: Validate after user fills the missing fields popup.
    Returns either more missing fields OR ready_to_draft=true + final_data.
    """
    draft_type = _normalize_type(req.draft_type)

    result = _ollama_validate(
        draft_type, req.original_input,
        req.user_responses, req.extracted_so_far,
    )

    if not result or "all_complete" not in result:
        result = _fallback_validate(draft_type, req.user_responses, req.extracted_so_far)
        result["source"] = "rule_based"
    else:
        result["source"] = "ollama"

    return {
        "success":            True,
        "still_missing":      result.get("still_missing", []),
        "all_complete":       result.get("all_complete", False),
        "completeness_score": result.get("completeness_score", 0),
        "issues_found":       result.get("issues_found", []),
        "ready_to_draft":     result.get("ready_to_draft", False),
        "final_data":         result.get("final_data", {}),
        "show_popup":         len(result.get("still_missing", [])) > 0,
        "source":             result.get("source", "unknown"),
    }


@router.post("/drafter/step3-generate")
def step3_generate(req: Step3Request):
    """
    Step 3: Generate the complete petition draft using master prompts.
    Only called when ready_to_draft = True.
    """
    if not req.final_data:
        raise HTTPException(400, "No case data provided")

    draft_type = _normalize_type(req.draft_type)

    rag_ctx = req.rag_context or _get_rag_context(
        f"{draft_type} {req.final_data.get('ground', '')} divorce cruelty Hindu Marriage Act India",
        top_k=5,
    )

    draft  = _ollama_generate(draft_type, req.final_data, rag_ctx,
                               side=req.side or "petitioner",
                               petition_text=req.petition_text or "")
    source = "ollama_llm"

    # Fallback: template
    if not draft or len(draft.strip()) < 300:
        draft  = _template_generate(draft_type, req.final_data)
        source = "rule_based_template"

    return {
        "success":    True,
        "draft":      draft,
        "draft_type": draft_type,
        "case_data":  req.final_data,
        "word_count": len(draft.split()),
        "source":     source,
    }


@router.post("/drafter/export-docx")
def export_docx(req: DocxExportRequest):
    """
    Export the draft text as a formatted Word (.docx) document.
    Bold formatting applied to ALL-CAPS lines, headers, party labels.
    """
    if not req.draft_text.strip():
        raise HTTPException(400, "No draft text provided")

    try:
        docx_bytes = _build_docx(req.draft_text)
    except Exception as e:
        raise HTTPException(500, f"DOCX generation failed: {str(e)}")

    filename = (req.filename or "LegalOne_Draft").replace(" ", "_") + ".docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/drafter/required-fields/{draft_type}")
def get_required_fields(draft_type: str):
    """Returns field schema for a draft type — used to build the popup UI."""
    normalized = _normalize_type(draft_type)
    fields = REQUIRED_FIELDS.get(normalized)
    if not fields:
        raise HTTPException(404, f"Unknown draft type: {draft_type}")
    return {
        "draft_type": normalized,
        "fields": [
            {
                "key":      k,
                "label":    v["label"],
                "required": v["required"],
                "question": v["ask"],
            }
            for k, v in fields.items()
        ],
    }
