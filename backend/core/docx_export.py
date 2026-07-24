
"""
docx_export.py — legalone/backend/core/docx_export.py
Generates a .docx matching the EXACT format of the uploaded sample:
  - Bold headings (court name, case no, party labels, section headings)
  - Normal body text
  - Right-aligned party labels (…PETITIONER / …RESPONDENT)
  - Centered VERSUS
  - Proper spacing between sections
"""

import io
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
from docx import Document
from docx.shared import Pt, Inches


# Lines that should be rendered BOLD (matches sample doc exactly)
BOLD_LINE_PATTERNS = [
    r'^IN THE COURT OF',
    r'^MATRIMONIAL CASE NO\.',
    r'^PETITION FOR DECREE',
    r'^WRITTEN STATEMENT',
    r'^WRITTEN ARGUMENTS',
    r'^COUNTER\b',
    r'^MOST RESPECTFULLY SHOWETH',
    r'^P\s*R\s*A\s*Y\s*E\s*R$',
    r'^PRAYER$',
    r'^VERIFICATION$',
    r'^PRELIMINARY OBJECTIONS',
    r'^CASE FACTS',
    r'^PARAWISE REPLY',
    r'^REPLY TO PRAYER',
    r'^WITHOUT PREJUDICE',
    r'^IN THE MATTER OF',
    r'^THROUGH$',
    r'^VERSUS$',
]

# Lines that end with these should be right-aligned (party role labels)
RIGHT_ALIGN_SUFFIXES = ['PETITIONER', 'RESPONDENT', 'APPELLANT', 'DEFENDANT']

CENTER_KEYWORDS = ['VERSUS']


def _is_bold_line(line: str) -> bool:
    stripped = line.strip()
    return any(re.match(pat, stripped, re.IGNORECASE) for pat in BOLD_LINE_PATTERNS)


def _is_party_role_line(line: str) -> bool:
    stripped = line.strip()
    return any(stripped.endswith(suf) for suf in RIGHT_ALIGN_SUFFIXES) and '...' in stripped


def _is_center_line(line: str) -> bool:
    return line.strip().upper() in CENTER_KEYWORDS


def _is_numbered_para(line: str) -> bool:
    return bool(re.match(r'^\d+\.\s+\S', line.strip()))


def _is_prayer_item(line: str) -> bool:
    return bool(re.match(r'^\([a-z]\)\s', line.strip()))


def generate_docx(draft_text: str, metadata: dict = None) -> bytes:
    """
    Convert plain draft text into a formatted .docx matching
    the exact style of the sample Indian court petition.
    """
    doc = Document()

    # ── Page setup — matches standard Indian legal pleading ──
    section = doc.sections[0]
    section.page_height  = Inches(11.69)   # A4
    section.page_width   = Inches(8.27)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Default style — Times New Roman 12pt ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(6)

    lines = draft_text.split('\n')

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        # Blank line → small spacer paragraph
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            continue

        para = doc.add_paragraph()

        # ── Alignment ──
        if _is_center_line(stripped):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif _is_party_role_line(stripped):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif _is_bold_line(stripped) and stripped.upper() in (
            'PRAYER', 'VERSUS', 'VERIFICATION', 'THROUGH', 'P R A Y E R'
        ):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ── Indentation for numbered paragraphs and prayer items ──
        if _is_numbered_para(stripped):
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.first_line_indent = Inches(-0.3)
        elif _is_prayer_item(stripped):
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)

        # ── Add run with bold detection ──
        run = para.add_run(stripped)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        if _is_bold_line(stripped):
            run.font.bold = True
            run.font.size = Pt(13) if any(
                k in stripped.upper() for k in ('IN THE COURT', 'PETITION FOR', 'WRITTEN STATEMENT', 'WRITTEN ARGUMENTS')
            ) else Pt(12)
        elif _is_party_role_line(stripped):
            run.font.bold = True

    # ── Save to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── FastAPI route helper ──────────────────────────────────────
ROUTE_SNIPPET = '''
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from core.docx_export import generate_docx

router = APIRouter()

class DocxRequest(BaseModel):
    draft_text: str

@router.post("/master-drafter/export-docx")
def export_docx(req: DocxRequest):
    if not req.draft_text.strip():
        raise HTTPException(400, "Draft text is empty")
    try:
        docx_bytes = generate_docx(req.draft_text)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="petition.docx"'}
        )
    except Exception as e:
        raise HTTPException(500, str(e))
'''

if __name__ == "__main__":
    sample = """IN THE COURT OF THE PRINCIPAL JUDGE, FAMILY COURT AT NEW DELHI

MATRIMONIAL CASE NO. __________ OF 2025

IN THE MATTER OF:

Mr. Rohan Malhotra, aged about 38 years,
Residing at B-124, Greater Kailash-I, New Delhi.                    ...PETITIONER

VERSUS

Mrs. Priya Malhotra, aged about 35 years,
Residing at 456, Lajpat Nagar, New Delhi.                            ...RESPONDENT

PETITION FOR DECREE OF DIVORCE UNDER SECTION 13(1)(ia) OF HINDU MARRIAGE ACT, 1955

MOST RESPECTFULLY SHOWETH:

1.  That the marriage of the Petitioner and Respondent was solemnized on 18.02.2016 at Rajouri Garden, New Delhi according to Hindu rites and ceremonies.

2.  That one male child was born out of the wedlock.

PRAYER

That this Hon'ble Court be pleased to:

(a)  Pass a decree of divorce dissolving the marriage;
(b)  Pass such other order as deemed fit.

VERIFICATION

I, Rohan Malhotra, verify that the contents are true."""

    docx_bytes = generate_docx(sample)
    with open('/mnt/user-data/outputs/test_petition_format.docx', 'wb') as f:
        f.write(docx_bytes)
    print(f"✓ Test DOCX generated: {len(docx_bytes)} bytes")
